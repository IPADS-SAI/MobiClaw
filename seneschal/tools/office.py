# -*- coding: utf-8 -*-
"""Office document tools (DOCX/PDF/XLSX)."""

from __future__ import annotations

import logging
import os
import textwrap
from pathlib import Path
from typing import Any

from agentscope.message import TextBlock
from agentscope.tool import ToolResponse

logger = logging.getLogger(__name__)


def _pick_reportlab_font() -> tuple[str, str]:
    """Pick a ReportLab font that can render Chinese text when possible.

    Returns:
        (font_name, font_source)
    """
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception:
        return "Helvetica", "builtin"

    configured_pdf = (os.environ.get("SENESCHAL_PDF_FONT_PATH") or "").strip()
    configured_cjk = (os.environ.get("SENESCHAL_CJK_FONT_PATH") or "").strip()
    candidates: list[str] = []
    if configured_pdf:
        candidates.append(configured_pdf)
    if configured_cjk and configured_cjk not in candidates:
        candidates.append(configured_cjk)
    candidates.extend(
        [
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.otf",
            "/usr/share/fonts/truetype/arphic/uming.ttc",
        ]
    )

    for path in candidates:
        candidate = Path(path).expanduser()
        if not candidate.exists() or not candidate.is_file():
            continue
        font_name = "SeneschalCJK"
        try:
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, str(candidate)))
            return font_name, str(candidate)
        except Exception as exc:
            logger.warning("office.pdf.font_register_failed path=%s error=%s", candidate, exc)

    try:
        cid_font = "STSong-Light"
        if cid_font not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(UnicodeCIDFont(cid_font))
        return cid_font, "cid:STSong-Light"
    except Exception as exc:
        logger.warning("office.pdf.cid_font_register_failed error=%s", exc)

    return "Helvetica", "builtin"


def _resolve_write_path(path: str) -> tuple[Path | None, str | None]:
    resolved_path = (path or "").strip()
    if not resolved_path:
        return None, "empty_path"

    root = (os.environ.get("SENESCHAL_FILE_WRITE_ROOT") or "").strip()
    target = Path(resolved_path).expanduser()
    if root:
        root_path = Path(root).expanduser().resolve()
        target = target if target.is_absolute() else root_path / target
        try:
            target = target.resolve()
        except FileNotFoundError:
            target = target.absolute()
        if target != root_path and root_path not in target.parents:
            return None, "path_outside_root"
    else:
        target = target.resolve()
    return target, None


def _trim_text(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    return text[: max_chars - 3].rstrip() + "..."


async def read_docx_text(
    file_path: str,
    max_chars: int | None = None,
) -> ToolResponse:
    """Read text content from a DOCX file.

    Args:
        file_path: Local DOCX file path.
        max_chars: Optional maximum number of returned characters.
    """
    resolved_path = (file_path or "").strip()
    if not resolved_path:
        return ToolResponse(
            content=[TextBlock(type="text", text="[DOCX] Empty file path.")],
            metadata={"error": "empty_path"},
        )

    target = Path(resolved_path).expanduser()
    if not target.exists():
        return ToolResponse(
            content=[TextBlock(type="text", text="[DOCX] File not found.")],
            metadata={"error": "not_found", "path": str(target)},
        )

    if target.suffix.lower() != ".docx":
        return ToolResponse(
            content=[TextBlock(type="text", text="[DOCX] Only .docx files are supported.")],
            metadata={"error": "invalid_extension", "path": str(target)},
        )

    try:
        from docx import Document
    except ImportError:
        return ToolResponse(
            content=[
                TextBlock(
                    type="text",
                    text="[DOCX] Missing dependency: install python-docx to read DOCX files.",
                )
            ],
            metadata={"error": "missing_dependency"},
        )

    max_chars_env = os.environ.get("SENESCHAL_DOCX_MAX_CHARS", "12000")
    max_chars_value = int(max_chars) if max_chars is not None else int(max_chars_env)
    max_chars_value = max(1000, min(max_chars_value, 200000))

    try:
        doc = Document(str(target))
    except Exception as exc:
        return ToolResponse(
            content=[TextBlock(type="text", text=f"[DOCX] Failed to read DOCX: {exc}")],
            metadata={"error": str(exc)},
        )

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    table_count = 0
    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text:
                parts.append(row_text)

    combined = "\n".join(parts)
    combined = _trim_text(combined, max_chars_value)

    return ToolResponse(
        content=[TextBlock(type="text", text=f"[DOCX] {target}\n{combined}")],
        metadata={
            "path": str(target),
            "paragraphs": len(doc.paragraphs),
            "tables": table_count,
        },
    )


async def create_docx_from_text(
    output_path: str,
    content: str,
    title: str | None = None,
) -> ToolResponse:
    """Create a DOCX file from plain text.

    Args:
        output_path: Output DOCX file path.
        content: Plain text content to write.
        title: Optional document title inserted as heading.
    """
    target, error = _resolve_write_path(output_path)
    if error == "empty_path":
        return ToolResponse(
            content=[TextBlock(type="text", text="[DOCX] Empty output path.")],
            metadata={"error": error},
        )
    if error == "path_outside_root":
        return ToolResponse(
            content=[TextBlock(type="text", text="[DOCX] Path is outside SENESCHAL_FILE_WRITE_ROOT.")],
            metadata={"error": error},
        )

    try:
        from docx import Document
    except ImportError:
        return ToolResponse(
            content=[
                TextBlock(
                    type="text",
                    text="[DOCX] Missing dependency: install python-docx to create DOCX files.",
                )
            ],
            metadata={"error": "missing_dependency"},
        )

    assert target is not None
    if target.suffix.lower() != ".docx":
        return ToolResponse(
            content=[TextBlock(type="text", text="[DOCX] Output must be a .docx file.")],
            metadata={"error": "invalid_extension", "path": str(target)},
        )
    try:
        doc = Document()
        if title:
            doc.add_heading(title, level=1)

        normalized = (content or "").replace("\r\n", "\n").replace("\r", "\n")
        blocks = normalized.split("\n\n") if normalized else [""]
        for block in blocks:
            paragraph = doc.add_paragraph()
            for idx, line in enumerate(block.splitlines()):
                run = paragraph.add_run(line)
                if idx < len(block.splitlines()) - 1:
                    run.add_break()

        target.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(target))
    except Exception as exc:
        return ToolResponse(
            content=[TextBlock(type="text", text=f"[DOCX] Write failed: {exc}")],
            metadata={"error": str(exc)},
        )

    logger.info("office.create_docx path=%s", target)
    return ToolResponse(
        content=[TextBlock(type="text", text=f"[DOCX] Wrote: {target}")],
        metadata={"path": str(target)},
    )


async def edit_docx(
    file_path: str,
    output_path: str,
    replacements: list[dict[str, str]] | None = None,
    append_paragraphs: list[str] | None = None,
    tables: list[list[list[str]]] | None = None,
) -> ToolResponse:
    """Edit a DOCX file with find/replace, appends, and simple tables.

    Note: paragraph-level replacements rewrite paragraph text and may drop formatting.

    Args:
        file_path: Source DOCX file path.
        output_path: Output DOCX file path.
        replacements: Optional replacement pairs, each item uses keys ``old`` and ``new``.
        append_paragraphs: Optional paragraphs appended to the end of document.
        tables: Optional table data to append, each table is a 2D list of cell values.
    """
    resolved_path = (file_path or "").strip()
    if not resolved_path:
        return ToolResponse(
            content=[TextBlock(type="text", text="[DOCX] Empty input path.")],
            metadata={"error": "empty_path"},
        )

    source = Path(resolved_path).expanduser()
    if not source.exists():
        return ToolResponse(
            content=[TextBlock(type="text", text="[DOCX] File not found.")],
            metadata={"error": "not_found", "path": str(source)},
        )

    target, error = _resolve_write_path(output_path)
    if error == "empty_path":
        return ToolResponse(
            content=[TextBlock(type="text", text="[DOCX] Empty output path.")],
            metadata={"error": error},
        )
    if error == "path_outside_root":
        return ToolResponse(
            content=[TextBlock(type="text", text="[DOCX] Path is outside SENESCHAL_FILE_WRITE_ROOT.")],
            metadata={"error": error},
        )

    try:
        from docx import Document
    except ImportError:
        return ToolResponse(
            content=[
                TextBlock(
                    type="text",
                    text="[DOCX] Missing dependency: install python-docx to edit DOCX files.",
                )
            ],
            metadata={"error": "missing_dependency"},
        )

    replace_pairs = []
    for item in replacements or []:
        old = (item or {}).get("old")
        new = (item or {}).get("new")
        if old is None or new is None:
            continue
        replace_pairs.append((str(old), str(new)))

    assert target is not None
    if target.suffix.lower() != ".docx":
        return ToolResponse(
            content=[TextBlock(type="text", text="[DOCX] Output must be a .docx file.")],
            metadata={"error": "invalid_extension", "path": str(target)},
        )
    try:
        doc = Document(str(source))
        replacement_count = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text or ""
            new_text = text
            for old, new in replace_pairs:
                if old in new_text:
                    new_text = new_text.replace(old, new)
            if new_text != text:
                paragraph.text = new_text
                replacement_count += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text or ""
                        new_text = text
                        for old, new in replace_pairs:
                            if old in new_text:
                                new_text = new_text.replace(old, new)
                        if new_text != text:
                            paragraph.text = new_text
                            replacement_count += 1

        for para in append_paragraphs or []:
            doc.add_paragraph(str(para))

        for table_data in tables or []:
            if not table_data:
                continue
            row_count = len(table_data)
            col_count = max(len(row) for row in table_data if row) if table_data else 0
            if row_count <= 0 or col_count <= 0:
                continue
            table = doc.add_table(rows=row_count, cols=col_count)
            for r_idx, row in enumerate(table_data):
                for c_idx, value in enumerate(row or []):
                    table.cell(r_idx, c_idx).text = str(value)

        target.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(target))
    except Exception as exc:
        return ToolResponse(
            content=[TextBlock(type="text", text=f"[DOCX] Edit failed: {exc}")],
            metadata={"error": str(exc)},
        )

    return ToolResponse(
        content=[TextBlock(type="text", text=f"[DOCX] Wrote: {target}")],
        metadata={
            "path": str(target),
            "replacements": len(replace_pairs),
            "replacement_paragraphs": replacement_count,
        },
    )


async def create_pdf_from_text(
    output_path: str,
    content: str,
    title: str | None = None,
    page_size: str = "letter",
) -> ToolResponse:
    """Create a PDF from plain text using reportlab.

    Args:
        output_path: Output PDF file path.
        content: Plain text content to render.
        title: Optional PDF title drawn at the top.
        page_size: Page size name, supports ``letter`` and ``a4``.
    """
    target, error = _resolve_write_path(output_path)
    if error == "empty_path":
        return ToolResponse(
            content=[TextBlock(type="text", text="[PDF] Empty output path.")],
            metadata={"error": error},
        )
    if error == "path_outside_root":
        return ToolResponse(
            content=[TextBlock(type="text", text="[PDF] Path is outside SENESCHAL_FILE_WRITE_ROOT.")],
            metadata={"error": error},
        )

    try:
        from reportlab.lib.pagesizes import A4, letter
        from reportlab.pdfgen import canvas
    except ImportError:
        return ToolResponse(
            content=[
                TextBlock(
                    type="text",
                    text="[PDF] Missing dependency: install reportlab to create PDF files.",
                )
            ],
            metadata={"error": "missing_dependency"},
        )

    pagesize = letter if (page_size or "letter").lower() == "letter" else A4
    assert target is not None
    if target.suffix.lower() != ".pdf":
        return ToolResponse(
            content=[TextBlock(type="text", text="[PDF] Output must be a .pdf file.")],
            metadata={"error": "invalid_extension", "path": str(target)},
        )

    try:
        target.parent.mkdir(parents=True, exist_ok=True)
        pdf = canvas.Canvas(str(target), pagesize=pagesize)
        width, height = pagesize
        margin = 72
        y = height - margin
        font_name, font_source = _pick_reportlab_font()
        logger.info("office.pdf.font_selected name=%s source=%s", font_name, font_source)
        if font_source == "builtin":
            logger.warning(
                "office.pdf.no_cjk_font_found; set SENESCHAL_PDF_FONT_PATH or SENESCHAL_CJK_FONT_PATH to a valid CJK font file to avoid garbled Chinese text"
            )
        pdf.setFont(font_name, 12)

        if title:
            pdf.setFont(font_name, 14)
            pdf.drawString(margin, y, title)
            y -= 24
            pdf.setFont(font_name, 12)

        wrapped_lines: list[str] = []
        for raw_line in (content or "").splitlines():
            if not raw_line.strip():
                wrapped_lines.append("")
                continue
            wrapped_lines.extend(textwrap.wrap(raw_line, width=100))

        for line in wrapped_lines:
            if y <= margin:
                pdf.showPage()
                pdf.setFont(font_name, 12)
                y = height - margin
            pdf.drawString(margin, y, line)
            y -= 14

        pdf.save()
    except Exception as exc:
        return ToolResponse(
            content=[TextBlock(type="text", text=f"[PDF] Write failed: {exc}")],
            metadata={"error": str(exc)},
        )

    return ToolResponse(
        content=[TextBlock(type="text", text=f"[PDF] Wrote: {target}")],
        metadata={"path": str(target), "font_name": font_name, "font_source": font_source},
    )


async def read_xlsx_summary(
    file_path: str,
    max_rows: int | None = 20,
) -> ToolResponse:
    """Read an XLSX file and return a summary plus a small preview.

    Args:
        file_path: Local XLSX or XLSM file path.
        max_rows: Maximum preview rows per sheet.
    """
    resolved_path = (file_path or "").strip()
    if not resolved_path:
        return ToolResponse(
            content=[TextBlock(type="text", text="[XLSX] Empty file path.")],
            metadata={"error": "empty_path"},
        )

    target = Path(resolved_path).expanduser()
    if not target.exists():
        return ToolResponse(
            content=[TextBlock(type="text", text="[XLSX] File not found.")],
            metadata={"error": "not_found", "path": str(target)},
        )

    if target.suffix.lower() not in {".xlsx", ".xlsm"}:
        return ToolResponse(
            content=[TextBlock(type="text", text="[XLSX] Only .xlsx/.xlsm files are supported.")],
            metadata={"error": "invalid_extension", "path": str(target)},
        )

    try:
        import pandas as pd
        from openpyxl import load_workbook
    except ImportError as exc:
        return ToolResponse(
            content=[TextBlock(type="text", text=f"[XLSX] Missing dependency: {exc}")],
            metadata={"error": "missing_dependency"},
        )

    preview_rows = max(1, min(int(max_rows or 20), 200))

    try:
        wb = load_workbook(filename=str(target), read_only=True, data_only=True)
    except Exception as exc:
        return ToolResponse(
            content=[TextBlock(type="text", text=f"[XLSX] Failed to read workbook: {exc}")],
            metadata={"error": str(exc)},
        )

    summary: list[dict[str, Any]] = []
    lines = [f"[XLSX] {target}"]
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        max_row = sheet.max_row or 0
        max_col = sheet.max_column or 0
        preview = []
        try:
            df = pd.read_excel(str(target), sheet_name=sheet_name, nrows=preview_rows)
            preview = df.to_dict(orient="records")
        except Exception:
            preview = []

        summary.append(
            {
                "name": sheet_name,
                "rows": max_row,
                "cols": max_col,
                "preview": preview,
            }
        )
        lines.append(f"- {sheet_name}: rows={max_row}, cols={max_col}, preview_rows={len(preview)}")

    wb.close()

    return ToolResponse(
        content=[TextBlock(type="text", text="\n".join(lines))],
        metadata={"path": str(target), "sheets": summary},
    )


async def write_xlsx_from_records(
    output_path: str,
    records: list[dict[str, Any]] | None = None,
    sheet_name: str = "Sheet1",
) -> ToolResponse:
    """Write an XLSX file from a list of dictionaries.

    Args:
        output_path: Output XLSX or XLSM file path.
        records: List of dict records, each dict maps column names to values.
        sheet_name: Target worksheet name.
    """
    target, error = _resolve_write_path(output_path)
    if error == "empty_path":
        return ToolResponse(
            content=[TextBlock(type="text", text="[XLSX] Empty output path.")],
            metadata={"error": error},
        )
    if error == "path_outside_root":
        return ToolResponse(
            content=[TextBlock(type="text", text="[XLSX] Path is outside SENESCHAL_FILE_WRITE_ROOT.")],
            metadata={"error": error},
        )

    try:
        import pandas as pd
    except ImportError:
        return ToolResponse(
            content=[TextBlock(type="text", text="[XLSX] Missing dependency: install pandas.")],
            metadata={"error": "missing_dependency"},
        )

    assert target is not None
    if target.suffix.lower() not in {".xlsx", ".xlsm"}:
        return ToolResponse(
            content=[TextBlock(type="text", text="[XLSX] Output must be a .xlsx/.xlsm file.")],
            metadata={"error": "invalid_extension", "path": str(target)},
        )
    try:
        df = pd.DataFrame.from_records(records or [])
        target.parent.mkdir(parents=True, exist_ok=True)
        df.to_excel(str(target), index=False, sheet_name=sheet_name)
    except Exception as exc:
        return ToolResponse(
            content=[TextBlock(type="text", text=f"[XLSX] Write failed: {exc}")],
            metadata={"error": str(exc)},
        )

    return ToolResponse(
        content=[TextBlock(type="text", text=f"[XLSX] Wrote: {target}")],
        metadata={"path": str(target), "rows": len(records or [])},
    )


async def write_xlsx_from_rows(
    output_path: str,
    rows: list[list[Any]] | None = None,
    headers: list[str] | None = None,
    sheet_name: str = "Sheet1",
) -> ToolResponse:
    """Write an XLSX file from row data.

    Args:
        output_path: Output XLSX or XLSM file path.
        rows: Row-major table data.
        headers: Optional column headers.
        sheet_name: Target worksheet name.
    """
    target, error = _resolve_write_path(output_path)
    if error == "empty_path":
        return ToolResponse(
            content=[TextBlock(type="text", text="[XLSX] Empty output path.")],
            metadata={"error": error},
        )
    if error == "path_outside_root":
        return ToolResponse(
            content=[TextBlock(type="text", text="[XLSX] Path is outside SENESCHAL_FILE_WRITE_ROOT.")],
            metadata={"error": error},
        )

    try:
        import pandas as pd
    except ImportError:
        return ToolResponse(
            content=[TextBlock(type="text", text="[XLSX] Missing dependency: install pandas.")],
            metadata={"error": "missing_dependency"},
        )

    assert target is not None
    if target.suffix.lower() not in {".xlsx", ".xlsm"}:
        return ToolResponse(
            content=[TextBlock(type="text", text="[XLSX] Output must be a .xlsx/.xlsm file.")],
            metadata={"error": "invalid_extension", "path": str(target)},
        )
    try:
        df = pd.DataFrame(rows or [], columns=headers)
        target.parent.mkdir(parents=True, exist_ok=True)
        df.to_excel(str(target), index=False, sheet_name=sheet_name)
    except Exception as exc:
        return ToolResponse(
            content=[TextBlock(type="text", text=f"[XLSX] Write failed: {exc}")],
            metadata={"error": str(exc)},
        )

    return ToolResponse(
        content=[TextBlock(type="text", text=f"[XLSX] Wrote: {target}")],
        metadata={"path": str(target), "rows": len(rows or [])},
    )
